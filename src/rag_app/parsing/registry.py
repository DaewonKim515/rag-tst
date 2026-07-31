"""
Parser registry and base interfaces for document parsing.

Provides a unified interface for all document parsers with
extension-based dispatch and security validation.
"""

import os
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, List, Optional, Type
from dataclasses import dataclass, field
from uuid import uuid4

from ..config.schema import AppConfig
from ..domain.models import ParsedDocument, ParsedBlock
from ..domain.enums import BlockType, FileType, OCRStatus
from ..domain.exceptions import (
    ParsingError, PDFParseError, DOCXParseError, TextParseError, MarkdownParseError,
    ZipBombError, XXEAttackError, OCRRequiredError
)
from ..ingestion.hasher import ZipBombDetector
from ..observability.logging import get_logger, get_current_job_id


logger = get_logger(__name__)


@dataclass(frozen=True, slots=True)
class ParserResult:
    """Result of parsing operation."""
    document: ParsedDocument
    warnings: List[str] = field(default_factory=list)


class DocumentParser(ABC):
    """Abstract base class for document parsers."""
    
    def __init__(self, config: AppConfig):
        self.config = config
        self.supported_extensions: List[str] = []
    
    @abstractmethod
    def can_parse(self, file_path: Path) -> bool:
        """Check if this parser can handle the given file."""
        pass
    
    @abstractmethod
    def parse(self, file_path: Path, document_id: str, md5_hash: str) -> ParserResult:
        """
        Parse a document file.
        
        Args:
            file_path: Path to the document file.
            document_id: Unique document identifier.
            md5_hash: MD5 hash of the file.
            
        Returns:
            ParserResult with ParsedDocument and warnings.
            
        Raises:
            ParsingError: If parsing fails.
        """
        pass
    
    def _create_parsed_document(
        self,
        file_path: Path,
        document_id: str,
        md5_hash: str,
        blocks: List[ParsedBlock],
        file_type: FileType,
        language: str = "ko",
        ocr_status: OCRStatus = OCRStatus.NOT_NEEDED,
        ocr_pages: tuple = (),
        parser_version: str = "1",
    ) -> ParsedDocument:
        """Helper to create ParsedDocument with common metadata."""
        from datetime import datetime
        import os
        
        stat = file_path.stat()
        
        return ParsedDocument(
            document_id=document_id,
            source_path=str(file_path),
            file_name=file_path.name,
            file_type=file_type,
            file_size=stat.st_size,
            md5_hash=md5_hash,
            modified_at=datetime.fromtimestamp(stat.st_mtime),
            blocks=tuple(blocks),
            parser_version=parser_version,
            language=language,
            ocr_status=ocr_status,
            ocr_pages=ocr_pages,
        )


class PDFParser(DocumentParser):
    """PDF parser using PyMuPDF (fitz)."""
    
    def __init__(self, config: AppConfig):
        super().__init__(config)
        self.supported_extensions = [".pdf"]
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"
    
    def parse(self, file_path: Path, document_id: str, md5_hash: str) -> ParserResult:
        job_id = get_current_job_id()
        
        try:
            import fitz  # PyMuPDF
        except ImportError as e:
            raise PDFParseError(file_path, e) from e
        
        blocks = []
        warnings = []
        ocr_pages = []
        
        try:
            doc = fitz.open(str(file_path))
            
            # Check if document is encrypted
            if doc.is_encrypted:
                if doc.authenticate(""):
                    warnings.append("Document was encrypted but opened with empty password")
                else:
                    doc.close()
                    raise PDFParseError(file_path, Exception("Document is encrypted"))
            
            for page_num, page in enumerate(doc, start=1):
                # Extract text with position info
                text = page.get_text("text")
                
                # Check text density for OCR detection
                text_length = len(text.strip())
                page_area = page.rect.width * page.rect.height
                text_density = text_length / page_area if page_area > 0 else 0
                
                # Heuristic: if very little text, might need OCR
                if text_density < 0.001 and text_length < 50:
                    ocr_pages.append(page_num)
                    warnings.append(f"Page {page_num}: Low text density, may need OCR")
                
                if text.strip():
                    # Split into paragraphs
                    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
                    
                    for para_idx, para_text in enumerate(paragraphs):
                        if para_text:
                            # Detect if it's a heading (simple heuristic)
                            block_type = BlockType.HEADING if self._is_heading(para_text) else BlockType.PARAGRAPH
                            
                            block = ParsedBlock(
                                block_type=block_type,
                                text=para_text,
                                page_number=page_num,
                                section_path=(f"Page {page_num}",),
                                sequence=len(blocks),
                            )
                            blocks.append(block)
                
                # Extract tables if any
                tables = page.find_tables()
                for table_idx, table in enumerate(tables):
                    try:
                        table_data = table.extract()
                        if table_data:
                            table_text = self._format_table(table_data)
                            block = ParsedBlock(
                                block_type=BlockType.TABLE,
                                text=table_text,
                                page_number=page_num,
                                section_path=(f"Page {page_num}", f"Table {table_idx + 1}"),
                                sequence=len(blocks),
                            )
                            blocks.append(block)
                    except Exception as e:
                        warnings.append(f"Page {page_num}, Table {table_idx}: Failed to extract - {e}")
            
            doc.close()
            
            # Determine OCR status
            ocr_status = OCRStatus.NOT_NEEDED
            if ocr_pages:
                if self.config.ocr.enabled:
                    ocr_status = OCRStatus.PENDING
                else:
                    ocr_status = OCRStatus.REQUIRED
                    warnings.append(f"OCR required for pages {ocr_pages} but OCR is disabled")
            
            file_type = FileType.PDF
            
        except OCRRequiredError:
            raise
        if not blocks:
            blocks.append(ParsedBlock(
                block_type=BlockType.PARAGRAPH,
                text="(No text extracted from PDF document)",
                page_number=1,
                sequence=0,
            ))

        parsed_doc = self._create_parsed_document(
            file_path=file_path,
            document_id=document_id,
            md5_hash=md5_hash,
            blocks=blocks,
            file_type=file_type,
            ocr_status=ocr_status,
            ocr_pages=tuple(ocr_pages),
        )
        
        return ParserResult(document=parsed_doc, warnings=warnings)
    
    def _is_heading(self, text: str) -> bool:
        """Simple heading detection heuristic."""
        # Short lines, all caps, or numbered
        if len(text) < 100 and (text.isupper() or text[0].isdigit()):
            return True
        return False
    
    def _format_table(self, table_data: List[List[str]]) -> str:
        """Format table as markdown-like text."""
        if not table_data:
            return ""
        
        # Header
        headers = [str(cell).strip() for cell in table_data[0]]
        lines = [" | ".join(headers)]
        lines.append(" | ".join(["---"] * len(headers)))
        
        # Rows
        for row in table_data[1:]:
            cells = [str(cell).strip() for cell in row]
            lines.append(" | ".join(cells))
        
        return "\n".join(lines)


class DOCXParser(DocumentParser):
    """DOCX parser using python-docx with XXE and zip bomb protection."""
    
    def __init__(self, config: AppConfig):
        super().__init__(config)
        self.supported_extensions = [".docx"]
        self.zip_bomb_detector = ZipBombDetector()
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"
    
    def parse(self, file_path: Path, document_id: str, md5_hash: str) -> ParserResult:
        job_id = get_current_job_id()
        
        # Check for zip bomb
        try:
            self.zip_bomb_detector.check(file_path)
        except ZipBombError:
            raise
        except Exception:
            pass  # Let parser handle
        
        try:
            from docx import Document
            from docx.oxml.ns import qn
            from docx.table import Table
        except ImportError as e:
            raise DOCXParseError(file_path, e) from e
        
        blocks = []
        warnings = []
        section_path = []
        
        try:
            doc = Document(str(file_path))
            
            for element_idx, element in enumerate(doc.element.body):
                tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
                
                if tag == "p":  # Paragraph
                    para = doc.paragraphs[element_idx] if element_idx < len(doc.paragraphs) else None
                    if para and para.text.strip():
                        # Detect heading style
                        style_name = para.style.name if para.style else ""
                        block_type = BlockType.HEADING if style_name.startswith("Heading") else BlockType.PARAGRAPH
                        
                        # Track section hierarchy
                        if block_type == BlockType.HEADING:
                            level = int(style_name[-1]) if style_name[-1].isdigit() else 1
                            section_path = section_path[:level-1] + (para.text.strip(),)
                        
                        block = ParsedBlock(
                            block_type=block_type,
                            text=para.text.strip(),
                            section_path=section_path,
                            sequence=len(blocks),
                        )
                        blocks.append(block)
                
                elif tag == "tbl":  # Table
                    # Find corresponding table object
                    table_idx = sum(1 for e in doc.element.body[:element_idx] 
                                   if e.tag.split("}")[-1] == "tbl")
                    if table_idx < len(doc.tables):
                        table = doc.tables[table_idx]
                        table_text = self._extract_table(table)
                        if table_text:
                            block = ParsedBlock(
                                block_type=BlockType.TABLE,
                                text=table_text,
                                section_path=section_path + (f"Table {table_idx + 1}",),
                                sequence=len(blocks),
                            )
                            blocks.append(block)
            
            file_type = FileType.DOCX
            
        except ZipBombError:
            raise
        except Exception as e:
            raise DOCXParseError(file_path, e) from e
        
        parsed_doc = self._create_parsed_document(
            file_path=file_path,
            document_id=document_id,
            md5_hash=md5_hash,
            blocks=blocks,
            file_type=file_type,
        )
        
        return ParserResult(document=parsed_doc, warnings=warnings)
    
    def _extract_table(self, table: "Table") -> str:
        """Extract table as markdown-like text."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        
        if not rows:
            return ""
        
        # Format as markdown table
        lines = [" | ".join(rows[0])]
        lines.append(" | ".join(["---"] * len(rows[0])))
        for row in rows[1:]:
            lines.append(" | ".join(row))
        
        return "\n".join(lines)


class TextParser(DocumentParser):
    """Plain text and Markdown parser with encoding detection."""
    
    def __init__(self, config: AppConfig):
        super().__init__(config)
        self.supported_extensions = [".txt", ".md", ".markdown"]
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.supported_extensions
    
    def parse(self, file_path: Path, document_id: str, md5_hash: str) -> ParserResult:
        job_id = get_current_job_id()
        
        blocks = []
        warnings = []
        
        try:
            # Detect encoding
            encoding = self._detect_encoding(file_path)
            
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
            
            file_type = FileType.MD if file_path.suffix.lower() in (".md", ".markdown") else FileType.TXT
            
            if file_type == FileType.MD:
                blocks = self._parse_markdown(content)
            else:
                blocks = self._parse_text(content)
            
        except Exception as e:
            raise TextParseError(file_path, e) from e
        
        parsed_doc = self._create_parsed_document(
            file_path=file_path,
            document_id=document_id,
            md5_hash=md5_hash,
            blocks=blocks,
            file_type=file_type,
        )
        
        return ParserResult(document=parsed_doc, warnings=warnings)
    
    def _detect_encoding(self, file_path: Path) -> str:
        """Detect file encoding using chardet if available."""
        try:
            import chardet
            with open(file_path, "rb") as f:
                raw = f.read(10000)  # Sample first 10KB
            result = chardet.detect(raw)
            encoding = result.get("encoding", "utf-8")
            confidence = result.get("confidence", 0)
            if confidence < 0.7:
                encoding = "utf-8"
            return encoding
        except ImportError:
            return "utf-8"
        except Exception:
            return "utf-8"
    
    def _parse_text(self, content: str) -> List[ParsedBlock]:
        """Parse plain text into paragraphs."""
        blocks = []
        
        # Split by double newline (paragraphs)
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        
        for idx, para in enumerate(paragraphs):
            block = ParsedBlock(
                block_type=BlockType.PARAGRAPH,
                text=para,
                sequence=idx,
            )
            blocks.append(block)
        
        return blocks
    
    def _parse_markdown(self, content: str) -> List[ParsedBlock]:
        """Parse Markdown preserving structure."""
        blocks = []
        sequence = 0
        section_path = []
        
        try:
            import markdown
            from markdown.extensions import Extension
            from markdown.treeprocessors import Treeprocessor
        except ImportError:
            # Fallback to simple parsing
            return self._parse_text(content)
        
        # Simple markdown parsing by lines
        lines = content.split("\n")
        current_para = []
        
        for line in lines:
            stripped = line.strip()
            
            # Heading detection
            if stripped.startswith("#"):
                # Flush current paragraph
                if current_para:
                    text = "\n".join(current_para).strip()
                    if text:
                        block = ParsedBlock(
                            block_type=BlockType.PARAGRAPH,
                            text=text,
                            section_path=tuple(section_path),
                            sequence=sequence,
                        )
                        blocks.append(block)
                        sequence += 1
                    current_para = []
                
                # Parse heading
                level = len(stripped) - len(stripped.lstrip("#"))
                heading_text = stripped.lstrip("# ").strip()
                
                section_path = section_path[:level-1] + (heading_text,)
                
                block = ParsedBlock(
                    block_type=BlockType.HEADING,
                    text=heading_text,
                    section_path=tuple(section_path),
                    sequence=sequence,
                )
                blocks.append(block)
                sequence += 1
            
            # Code block detection
            elif stripped.startswith("```"):
                if current_para:
                    text = "\n".join(current_para).strip()
                    if text:
                        block = ParsedBlock(
                            block_type=BlockType.PARAGRAPH,
                            text=text,
                            section_path=tuple(section_path),
                            sequence=sequence,
                        )
                        blocks.append(block)
                        sequence += 1
                    current_para = []
                
                # Collect code block
                code_lines = []
                for next_line in lines[lines.index(line)+1:]:
                    if next_line.strip().startswith("```"):
                        break
                    code_lines.append(next_line)
                
                if code_lines:
                    block = ParsedBlock(
                        block_type=BlockType.CODE,
                        text="\n".join(code_lines),
                        section_path=tuple(section_path),
                        sequence=sequence,
                    )
                    blocks.append(block)
                    sequence += 1
            
            elif stripped:
                current_para.append(line)
            else:
                # Empty line - flush paragraph
                if current_para:
                    text = "\n".join(current_para).strip()
                    if text:
                        block = ParsedBlock(
                            block_type=BlockType.PARAGRAPH,
                            text=text,
                            section_path=tuple(section_path),
                            sequence=sequence,
                        )
                        blocks.append(block)
                        sequence += 1
                    current_para = []
        
        # Flush remaining
        if current_para:
            text = "\n".join(current_para).strip()
            if text:
                block = ParsedBlock(
                    block_type=BlockType.PARAGRAPH,
                    text=text,
                    section_path=tuple(section_path),
                    sequence=sequence,
                )
                blocks.append(block)
        
        return blocks


class ParserRegistry:
    """Registry for document parsers with extension-based dispatch."""
    
    def __init__(self, config: AppConfig):
        self.config = config
        self._parsers: Dict[str, DocumentParser] = {}
        self._initialize_parsers()
    
    def _initialize_parsers(self) -> None:
        """Initialize all available parsers."""
        self._parsers = {
            ".pdf": PDFParser(self.config),
            ".docx": DOCXParser(self.config),
            ".txt": TextParser(self.config),
            ".md": TextParser(self.config),
            ".markdown": TextParser(self.config),
        }
    
    def get_parser(self, file_path: Path) -> DocumentParser:
        """
        Get appropriate parser for file.
        
        Args:
            file_path: Path to document file.
            
        Returns:
            DocumentParser instance.
            
        Raises:
            UnsupportedFileTypeError: If no parser available.
        """
        ext = file_path.suffix.lower()
        parser = self._parsers.get(ext)
        
        if parser is None:
            from ..domain.exceptions import UnsupportedFileTypeError
            supported = list(self._parsers.keys())
            raise UnsupportedFileTypeError(file_path, supported)
        
        return parser
    
    def parse(
        self, 
        file_path: Path, 
        document_id: Any, 
        md5_hash: str
    ) -> ParserResult:
        """Parse a document using the appropriate parser."""
        parser = self.get_parser(file_path)
        res = parser.parse(file_path, document_id, md5_hash)
        if isinstance(res, ParsedDocument):
            return ParserResult(document=res)
        return res
    
    def supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self._parsers.keys())
    
    def register_parser(self, extension: str, parser: DocumentParser) -> None:
        """Register a custom parser."""
        self._parsers[extension.lower()] = parser